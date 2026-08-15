"""What a job does to a document, and the processor contract that defines it.

``DocumentProcessor`` is what ``app.jobs.runJob`` executes; this module owns
that contract because the contract is about documents, not about how a job is
scheduled or tracked (that's ``app.jobManager``). ``DocumentAnalyzerProcessor``
is the real implementation: a job's ``documentLink`` is downloaded, its format
is detected, and metadata is pulled out of it. A ``.zip`` or ``.rar`` is
unpacked and every supported member analyzed the same way, so its
``DocumentMetadata.sourceKind`` is ``"folder"``; anything else is ``"single"``.

Unlike zip, rar has no format-decoding logic in Python's standard library --
``rarfile`` shells out to an ``unrar`` (or ``bsdtar``/``7z``) binary that must
already be on the machine. If it isn't, opening a ``.rar`` raises a clear
error rather than the library's own exception, which is otherwise easy to
mistake for a corrupt archive.

What "processing" means beyond metadata extraction -- chunking, embedding,
writing into a ``ragDbId`` -- is still undecided; this module only answers
"what did we just download".
"""

from __future__ import annotations

import csv
import io
import logging
import os
import zipfile
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import PurePosixPath
from typing import TYPE_CHECKING, Protocol
from urllib.parse import urlparse

import docx
import httpx
import pdfplumber
import rarfile
import tiktoken

# Where to find the unrar/bsdtar/7z binary rarfile shells out to. Not on every
# machine by default the way zip support is -- override per-environment with
# RAG_UNRAR_TOOL if it isn't on PATH under its usual name.
rarfile.UNRAR_TOOL = os.environ.get("RAG_UNRAR_TOOL", rarfile.UNRAR_TOOL)

if TYPE_CHECKING:
    # Only needed for the type hint below -- importing it for real would make
    # app.jobs and app.documents import each other.
    from app.jobs import Job

logger = logging.getLogger(__name__)

DOWNLOAD_TIMEOUT_SECONDS = 60.0
# A generous ceiling for a personal project, not a production sizing decision.
MAX_DOWNLOAD_BYTES = 200 * 1024 * 1024

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv"}
ARCHIVE_EXTENSIONS = {".zip", ".rar"}

# Used when the URL doesn't carry a recognizable extension (e.g. a bare
# `/download?id=123`), keyed on the response's Content-Type.
CONTENT_TYPE_EXTENSIONS = {
    "application/pdf": ".pdf",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/csv": ".csv",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/zip": ".zip",
    "application/x-zip-compressed": ".zip",
    "application/vnd.rar": ".rar",
    "application/x-rar-compressed": ".rar",
}

# cl100k_base (GPT-3.5/GPT-4's encoding) is used as a general-purpose token
# count, not tied to whichever model ends up doing embedding/generation --
# it's a proxy for "roughly how much content is here", useful for chunking
# decisions regardless of what reads the chunks later.
TOKEN_ENCODING_NAME = "cl100k_base"


@lru_cache(maxsize=1)
def _tokenEncoding() -> tiktoken.Encoding:
    # tiktoken fetches this encoding's BPE data over the network the first
    # time any process asks for it, then caches the file on disk (by default
    # under the system temp dir, which may not survive a container restart).
    # A deployment with no outbound network on a cold start will hit this
    # every time -- see _countTokens, which degrades to `None` rather than
    # failing the whole file when that happens.
    return tiktoken.get_encoding(TOKEN_ENCODING_NAME)


def _countTokens(text: str) -> int | None:
    """Best-effort token count for ``text``.

    Returns ``None`` -- not 0 -- if the tokenizer couldn't run, so a file the
    tokenizer failed on is distinguishable from a genuinely empty file.
    ``lru_cache`` above doesn't cache exceptions, so a transient failure (e.g.
    no network on first use) is retried on the next file rather than sticking
    for the rest of the process.
    """
    if not text:
        return 0
    try:
        return len(_tokenEncoding().encode(text))
    except Exception:
        logger.warning("Token counting unavailable.", exc_info=True)
        return None


class DownloadError(Exception):
    """The document could not be fetched."""


class UnsupportedDocumentError(Exception):
    """The file -- or a member of an archive -- is not one of the supported
    formats."""


class ArchiveToolMissingError(Exception):
    """The archive needs an external tool (e.g. unrar) that isn't available."""


@dataclass
class FileMetadata:
    """What we could tell about one file."""

    filename: str
    extension: str
    sizeBytes: int
    tableCount: int = 0
    imageCount: int = 0
    pageCount: int | None = None  # PDF
    wordCount: int | None = None  # docx, txt, md
    lineCount: int | None = None  # txt, md
    rowCount: int | None = None  # csv
    columnCount: int | None = None  # csv
    # Tokens in the file's extracted text (cl100k_base). None if the
    # tokenizer couldn't run; not attempted at all for formats with no
    # meaningful text (currently: none -- every supported format has some).
    tokenCount: int | None = None
    # Set instead of raising when one file (typically inside a zip) fails to
    # parse, so one bad member doesn't lose the metadata for the rest.
    error: str | None = None


@dataclass
class DocumentMetadata:
    """What we could tell about the thing that was downloaded.

    ``pageCount``/``imageCount``/``tableCount``/``tokenCount`` are totals
    across every file -- including files nested arbitrarily deep inside
    archives-within-archives, not just the top level. Per-file detail is
    still in ``files``; these are the roll-up most callers actually want.
    """

    sourceUrl: str
    sourceKind: str  # "single" | "folder"
    contentType: str | None
    fileCount: int
    totalSizeBytes: int
    pageCount: int
    imageCount: int
    tableCount: int
    tokenCount: int
    files: list[FileMetadata] = field(default_factory=list)


def _extensionOf(name: str) -> str:
    return PurePosixPath(name).suffix.lower()


def _filenameFromUrl(url: str) -> str:
    name = PurePosixPath(urlparse(url).path).name
    return name or "document"


async def download(url: str) -> tuple[bytes, str, str | None]:
    """Fetch ``url``. Returns the body, a filename to analyze it as, and the
    server's Content-Type (used as a fallback when the URL has no extension)."""
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=DOWNLOAD_TIMEOUT_SECONDS) as client:
            async with client.stream("GET", url) as response:
                response.raise_for_status()
                body = bytearray()
                async for chunk in response.aiter_bytes():
                    body.extend(chunk)
                    if len(body) > MAX_DOWNLOAD_BYTES:
                        raise DownloadError(
                            f"Document exceeds the {MAX_DOWNLOAD_BYTES // (1024 * 1024)}MB limit."
                        )
                contentType = response.headers.get("content-type")
    except httpx.HTTPError as exc:
        raise DownloadError(f"Could not download '{url}': {exc}") from exc

    return bytes(body), _filenameFromUrl(url), contentType


def _resolveExtension(filename: str, contentType: str | None) -> str:
    extension = _extensionOf(filename)
    if extension in SUPPORTED_EXTENSIONS or extension in ARCHIVE_EXTENSIONS:
        return extension

    if contentType:
        mime = contentType.split(";")[0].strip().lower()
        if mime in CONTENT_TYPE_EXTENSIONS:
            return CONTENT_TYPE_EXTENSIONS[mime]

    return extension  # Unrecognized either way; caller raises.


def analyzeBytes(filename: str, data: bytes, *, extension: str | None = None) -> FileMetadata:
    """Extract metadata from one file's contents.

    A parse failure is recorded on ``FileMetadata.error`` rather than raised,
    so analyzing a zip full of files doesn't abort on the first bad one.
    """
    extension = extension if extension is not None else _extensionOf(filename)
    meta = FileMetadata(filename=filename, extension=extension, sizeBytes=len(data))

    if extension not in SUPPORTED_EXTENSIONS:
        meta.error = f"Unsupported file type '{extension or filename}'."
        return meta

    try:
        if extension == ".pdf":
            _analyzePdf(data, meta)
        elif extension == ".docx":
            _analyzeDocx(data, meta)
        elif extension == ".csv":
            _analyzeCsv(data, meta)
        else:  # .txt, .md
            _analyzeText(data, meta)
    except Exception as exc:
        logger.warning("Failed to analyze '%s': %s", filename, exc)
        meta.error = str(exc)

    return meta


def _analyzePdf(data: bytes, meta: FileMetadata) -> None:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        meta.pageCount = len(pdf.pages)
        pageTexts = []
        for page in pdf.pages:
            meta.tableCount += len(page.find_tables())
            meta.imageCount += len(page.images)
            pageTexts.append(page.extract_text() or "")
    meta.tokenCount = _countTokens("\n".join(pageTexts))


def _analyzeDocx(data: bytes, meta: FileMetadata) -> None:
    document = docx.Document(io.BytesIO(data))
    meta.tableCount = len(document.tables)
    meta.imageCount = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    meta.wordCount = len(text.split())
    meta.tokenCount = _countTokens(text)


def _analyzeCsv(data: bytes, meta: FileMetadata) -> None:
    text = data.decode("utf-8-sig", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    meta.rowCount = len(rows)
    meta.columnCount = len(rows[0]) if rows else 0
    meta.tableCount = 1 if rows else 0
    meta.tokenCount = _countTokens(text)


def _analyzeText(data: bytes, meta: FileMetadata) -> None:
    text = data.decode("utf-8", errors="replace")
    meta.lineCount = text.count("\n") + (1 if text and not text.endswith("\n") else 0)
    meta.wordCount = len(text.split())
    meta.tokenCount = _countTokens(text)


# How many archive levels to unpack before giving up. Guards against a
# maliciously (or accidentally) deep chain of archives-within-archives eating
# unbounded memory/CPU -- five is far beyond anything a real submission needs.
# for now onlt single depth maybe in future we can increase it
MAX_ARCHIVE_DEPTH = 1


def _openArchive(extension: str, data: bytes):
    """Open a zip or rar so its members can be listed and read.

    ``zipfile.ZipFile`` and ``rarfile.RarFile`` expose the same
    ``infolist()`` / ``read(info)`` / ``is_dir()`` shape by design, so one
    analysis loop below handles either.
    """
    if extension == ".zip":
        return zipfile.ZipFile(io.BytesIO(data))
    return rarfile.RarFile(io.BytesIO(data))


def _analyzeArchiveEntries(
    extension: str, data: bytes, *, pathPrefix: str, depth: int
) -> list[FileMetadata]:
    """Walk one archive, recursing into any member that is itself a zip/rar.

    Returns a flat list -- a nested member's ``filename`` carries the full
    path (``outer.zip/inner.rar/doc.pdf``), so the hierarchy stays visible
    without the result needing to be tree-shaped.
    """
    if depth > MAX_ARCHIVE_DEPTH:
        return [
            FileMetadata(
                filename=pathPrefix.rstrip("/"),
                extension=extension,
                sizeBytes=len(data),
                error=f"Archive nesting exceeds the {MAX_ARCHIVE_DEPTH}-level limit.",
            )
        ]

    files: list[FileMetadata] = []
    try:
        with _openArchive(extension, data) as archive:
            for info in archive.infolist():
                name = PurePosixPath(info.filename).name
                # Skip directory entries and junk archives love to carry
                # (__MACOSX/, .DS_Store, Thumbs.db) rather than reporting
                # them as unsupported files.
                if info.is_dir() or not name or name.startswith("."):
                    continue

                memberPath = f"{pathPrefix}{info.filename}"
                memberExtension = _extensionOf(name)
                memberBytes = archive.read(info)

                if memberExtension in ARCHIVE_EXTENSIONS:
                    files.extend(
                        _analyzeArchiveEntries(
                            memberExtension,
                            memberBytes,
                            pathPrefix=f"{memberPath}/",
                            depth=depth + 1,
                        )
                    )
                else:
                    files.append(analyzeBytes(memberPath, memberBytes))
    except rarfile.RarExecError as exc:
        # Listing a rar's contents needs no external tool; reading a member's
        # bytes does. That means this can fail partway through the walk --
        # raised as one clear error instead of leaving `files` half-built.
        raise ArchiveToolMissingError(
            "Cannot read .rar archives: no unrar/bsdtar/7z tool found "
            f"(looked for '{rarfile.UNRAR_TOOL}'). Install one, or point "
            "RAG_UNRAR_TOOL at its path."
        ) from exc

    return files


def _analyzeArchive(
    url: str, extension: str, data: bytes, contentType: str | None
) -> DocumentMetadata:
    files = _analyzeArchiveEntries(extension, data, pathPrefix="", depth=0)

    return DocumentMetadata(
        sourceUrl=url,
        sourceKind="folder",
        contentType=contentType,
        fileCount=len(files),
        totalSizeBytes=sum(f.sizeBytes for f in files),
        pageCount=sum(f.pageCount or 0 for f in files),
        imageCount=sum(f.imageCount for f in files),
        tableCount=sum(f.tableCount for f in files),
        tokenCount=sum(f.tokenCount or 0 for f in files),
        files=files,
    )


def analyze(url: str, filename: str, data: bytes, contentType: str | None = None) -> DocumentMetadata:
    """Analyze a downloaded document. A zip or rar becomes a folder of files;
    anything else is a single file."""
    extension = _resolveExtension(filename, contentType)

    if extension in ARCHIVE_EXTENSIONS:
        return _analyzeArchive(url, extension, data, contentType)

    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(
            f"Unsupported document type for '{filename}' "
            f"(content-type: {contentType or 'unknown'})."
        )

    fileMeta = analyzeBytes(filename, data, extension=extension)
    return DocumentMetadata(
        sourceUrl=url,
        sourceKind="single",
        contentType=contentType,
        fileCount=1,
        totalSizeBytes=fileMeta.sizeBytes,
        pageCount=fileMeta.pageCount or 0,
        imageCount=fileMeta.imageCount,
        tableCount=fileMeta.tableCount,
        tokenCount=fileMeta.tokenCount or 0,
        files=[fileMeta],
    )


def extractText(filename: str, data: bytes, contentType: str | None = None) -> str:
    """The document's text, for callers that need the content rather than the
    counts ``analyze`` produces (chunking, embedding).

    Parses independently of ``analyze`` rather than sharing one pass: analyze
    walks a PDF page by page for tables and images and keeps only totals, so
    there is no extracted text left over to hand back. A document that is both
    analyzed and chunked is therefore parsed twice -- cheap enough at these
    sizes, and it keeps ``analyze``'s signature about metadata only.

    An archive becomes every supported member's text concatenated, in the
    order the archive lists them. A member that fails to parse is skipped with
    a warning, the same tolerance ``analyzeBytes`` has.
    """
    extension = _resolveExtension(filename, contentType)

    if extension in ARCHIVE_EXTENSIONS:
        return _archiveText(extension, data, depth=0)

    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(
            f"Cannot extract text from '{filename}' "
            f"(content-type: {contentType or 'unknown'})."
        )

    return _textOf(extension, data)


def _textOf(extension: str, data: bytes) -> str:
    if extension == ".pdf":
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if extension == ".docx":
        document = docx.Document(io.BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    if extension == ".csv":
        return data.decode("utf-8-sig", errors="replace")

    return data.decode("utf-8", errors="replace")  # .txt, .md


def _archiveText(extension: str, data: bytes, *, depth: int) -> str:
    if depth > MAX_ARCHIVE_DEPTH:
        return ""

    texts: list[str] = []
    try:
        with _openArchive(extension, data) as archive:
            for info in archive.infolist():
                name = PurePosixPath(info.filename).name
                if info.is_dir() or not name or name.startswith("."):
                    continue

                memberExtension = _extensionOf(name)
                memberBytes = archive.read(info)

                if memberExtension in ARCHIVE_EXTENSIONS:
                    texts.append(_archiveText(memberExtension, memberBytes, depth=depth + 1))
                elif memberExtension in SUPPORTED_EXTENSIONS:
                    try:
                        texts.append(_textOf(memberExtension, memberBytes))
                    except Exception as exc:
                        logger.warning("Skipping '%s': %s", info.filename, exc)
    except rarfile.RarExecError as exc:
        raise ArchiveToolMissingError(
            "Cannot read .rar archives: no unrar/bsdtar/7z tool found "
            f"(looked for '{rarfile.UNRAR_TOOL}'). Install one, or point "
            "RAG_UNRAR_TOOL at its path."
        ) from exc

    return "\n\n".join(text for text in texts if text.strip())


class DocumentProcessor(Protocol):
    """The one thing a job needs done to it.

    Whatever an implementation does -- fetch the link, chunk it, embed it,
    write it to a ragDbId -- it should raise on failure and otherwise mutate
    nothing but ``job.detail`` / ``job.metadata``; ``app.jobs.runJob`` owns
    ``job.status``.
    """

    async def process(self, job: "Job") -> None: ...


class StubDocumentProcessor:
    """Placeholder: marks every job done immediately, does no real work."""

    async def process(self, job: "Job") -> None:
        job.detail = "Processing is not implemented yet."


class DocumentAnalyzerProcessor:
    """Downloads ``job.documentLink``, analyzes it, and records the result.

    Any exception here -- a bad URL, an unsupported type, a network failure --
    is left to propagate; ``app.jobs.runJob`` is what catches it and marks
    the job failed.
    """

    async def process(self, job: "Job") -> None:
        data, filename, contentType = await download(job.documentLink)
        metadata = analyze(job.documentLink, filename, data, contentType)

        job.metadata = metadata
        job.detail = (
            f"Analyzed {metadata.fileCount} file(s) "
            f"({metadata.sourceKind}, {metadata.totalSizeBytes} bytes) -- "
            f"{metadata.pageCount} page(s), {metadata.imageCount} image(s), "
            f"{metadata.tableCount} table(s), {metadata.tokenCount} token(s)."
        )
